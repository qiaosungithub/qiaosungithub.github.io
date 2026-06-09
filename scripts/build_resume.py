from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "assets" / "pdf"
PDF_PATH = OUT_DIR / "Qiao-Sun-CV.pdf"
LEGACY_PDF_PATH = OUT_DIR / "resume.pdf"
DOCX_PATH = OUT_DIR / "Qiao-Sun-CV.docx"

BLUE = colors.HexColor("#1f4e79")
GRAY = colors.HexColor("#45515c")
LIGHT_GRAY = colors.HexColor("#d6dde3")


RESUME = {
    "name": "Qiao Sun",
    "contact": [
        "(+86) 17317700890",
        "(+1) 6172568632",
        "sqa24@mit.edu",
        "github.com/qiaosungithub",
    ],
    "sections": [
        {
            "title": "Research & Professional Experience",
            "entries": [
                {
                    "title": "Undergraduate Researcher, He Vision Group",
                    "meta": "Professor: Kaiming He",
                    "date": "Sep 2024 - Present",
                    "bullets": [
                        "Work as a UROP student in He Vision Group on generative models, including diffusion and flow matching, normalizing flows, and fast image generation.",
                        "Collaborate with Hanhong Zhao, Zhicheng Jiang, Xianbang Wang, and Yiyang Lu on a unified direction for vision-language understanding and text-to-image generation.",
                        "Conduct large-scale experiments with JAX on TPUs and PyTorch on GPUs; build training, evaluation, and ablation pipelines for image generation models.",
                        "Serve as one of the group leads and TPU/GCP managers/admins, helping coordinate compute resources, access, and shared training infrastructure.",
                    ],
                },
                {
                    "title": "Quant Strategy Analyst Intern, Ubiquant Investment (Jiukun Quant)",
                    "meta": "Quantitative research and strategy analysis",
                    "date": "Jun 20 - Aug 31, 2025",
                    "bullets": [
                        "Researched quantitative strategy ideas and backtesting workflows with Python-based data analysis.",
                        "Evaluated factor robustness, risk and turnover characteristics, and strategy diagnostics through statistical tests and visualization.",
                        "Built reproducible research utilities for experiment tracking, factor validation, and portfolio-performance analysis.",
                    ],
                },
            ],
        },
        {
            "title": "Selected Papers",
            "entries": [
                {
                    "title": "Is Noise Conditioning Necessary for Denoising Generative Models?",
                    "meta": "ICML 2025 poster; first author",
                    "date": "2025",
                    "bullets": [
                        "Revisited the necessity of noise conditioning in diffusion and flow-matching models by reimplementing eight denoising generative models across multiple datasets.",
                        "Proposed uEDM, a noise-unconditional diffusion model with competitive performance, supported by theory that matches empirical behavior.",
                    ],
                },
                {
                    "title": "Bidirectional Normalizing Flow: From Data to Noise and Back",
                    "meta": "CVPR Spotlight; project lead and first author",
                    "date": "2025",
                    "bullets": [
                        "Revisited normalizing flows with a learned reverse map, eliminating explicit inverse-flow computation and slow autoregressive inference.",
                        "Guided reverse learning through hidden-state alignment, enabling single-evaluation NF-based generation with state-of-the-art results among NF methods.",
                    ],
                },
                {
                    "title": "One-step Latent-free Image Generation with Pixel Mean Flows",
                    "meta": "ICML accepted; first author",
                    "date": "2026",
                    "bullets": [
                        "Built a strong baseline for one-step, pixel-space, latent-free generation using MeanFlow with x-prediction.",
                        "Reported 2.22 FID on ImageNet 256 and 2.48 FID on ImageNet 512; open-source implementation at Lyy-iiis/pMF.",
                    ],
                },
            ],
        },
        {
            "title": "Course Projects",
            "entries": [
                {
                    "title": "Fast Humanoid Loco-Manipulation via Flow Matching",
                    "meta": "MIT 6.4210 course project; robotics manipulation",
                    "date": "2025",
                    "bullets": [
                        "Project video: https://www.youtube.com/watch?v=MIpTpM4C71k&t=2s",
                        "Reimplemented a BeyondMimic-style humanoid loco-manipulation pipeline with simulated data preparation, diffusion training, and post-hoc control guidance.",
                        "Replaced DDPM sampling with flow matching and demonstrated lower-latency control with 5 sampling steps.",
                    ],
                },
            ],
        },
        {
            "title": "Education",
            "entries": [
                {
                    "title": "Massachusetts Institute of Technology",
                    "meta": "Undergraduate in Artificial Intelligence and Mathematics; expected graduation: 2028; GPA: 5.00/5.00",
                    "date": "Sep 2024 - Present",
                },
                {
                    "title": "Tsinghua University, Institute for Interdisciplinary Information Sciences",
                    "meta": "Pre-college student; GPA: 4.00/4.00",
                    "date": "Sep 2023 - Jul 2024",
                },
            ],
        },
        {
            "title": "Certificates & Honors",
            "page_break_before": True,
            "entries": [
                {
                    "title": "Top 17, 2025 Putnam Mathematical Competition",
                    "date": "2025",
                },
                {
                    "title": "2nd place, 2024 Putnam Mathematical Competition",
                    "date": "2024",
                },
                {
                    "title": "Gold Medal and 11th place, International Mathematical Olympiad",
                    "date": "2023",
                },
                {
                    "title": "Gold Medal and 1st place with perfect score, Chinese Mathematical Olympiad",
                    "date": "2022",
                },
                {
                    "title": "Excellent Award, Alibaba Global Mathematics Competition; top 70 among 50,000+ participants",
                    "date": "2022-2024",
                },
            ],
        },
        {
            "title": "Skills",
            "entries": [
                {
                    "title": "Programming",
                    "meta": "Python, C, PyTorch, JAX, GPU/TPU training and evaluation",
                },
                {
                    "title": "Languages",
                    "meta": "Mandarin (native), English (fluent)",
                },
            ],
        },
    ],
}


def register_fonts():
    font_dir = Path("/usr/share/fonts/truetype/dejavu")
    regular = font_dir / "DejaVuSans.ttf"
    bold = font_dir / "DejaVuSans-Bold.ttf"
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("ResumeSans", regular))
        pdfmetrics.registerFont(TTFont("ResumeSans-Bold", bold))
        return "ResumeSans", "ResumeSans-Bold"
    return "Helvetica", "Helvetica-Bold"


def make_pdf():
    base_font, bold_font = register_fonts()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=0.58 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.52 * inch,
        bottomMargin=0.56 * inch,
        title="Qiao Sun CV",
        author="Qiao Sun",
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Name",
            fontName=bold_font,
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Contact",
            fontName=base_font,
            fontSize=9,
            leading=11,
            textColor=GRAY,
            alignment=TA_CENTER,
            spaceAfter=9,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            fontName=bold_font,
            fontSize=10,
            leading=12,
            textColor=BLUE,
            spaceBefore=11,
            spaceAfter=3,
            uppercase=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryTitle",
            fontName=bold_font,
            fontSize=9.4,
            leading=11.5,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryDate",
            fontName=base_font,
            fontSize=8.7,
            leading=10.8,
            textColor=GRAY,
            alignment=TA_RIGHT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EntryMeta",
            fontName=base_font,
            fontSize=8.8,
            leading=10.8,
            textColor=GRAY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            fontName=base_font,
            fontSize=8.55,
            leading=10.8,
            leftIndent=12,
            firstLineIndent=-8,
            bulletIndent=0,
            spaceAfter=2.0,
        )
    )

    story = []
    story.append(Paragraph(escape(RESUME["name"]), styles["Name"]))
    story.append(Paragraph(escape(" / ".join(RESUME["contact"])), styles["Contact"]))
    story.append(HRFlowable(width="100%", thickness=0.7, color=LIGHT_GRAY, spaceAfter=2))

    width = A4[0] - doc.leftMargin - doc.rightMargin

    for section in RESUME["sections"]:
        if section.get("page_break_before") and story:
            story.append(PageBreak())
        story.append(Paragraph(escape(section["title"]).upper(), styles["Section"]))
        story.append(HRFlowable(width="100%", thickness=0.45, color=LIGHT_GRAY, spaceAfter=5))

        for entry in section["entries"]:
            title = Paragraph(escape(entry["title"]), styles["EntryTitle"])
            date = Paragraph(escape(entry.get("date", "")), styles["EntryDate"])
            block = [
                Table(
                    [[title, date]],
                    colWidths=[width * 0.76, width * 0.24],
                    style=TableStyle(
                        [
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                            ("TOPPADDING", (0, 0), (-1, -1), 0),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                        ]
                    ),
                )
            ]
            if entry.get("meta"):
                block.append(Paragraph(escape(entry["meta"]), styles["EntryMeta"]))
            for bullet in entry.get("bullets", []):
                block.append(Paragraph(f"<bullet>&bull;</bullet>{escape(bullet)}", styles["ResumeBullet"]))
            block.append(Spacer(1, 5))
            story.append(KeepTogether(block))

    doc.build(story)
    LEGACY_PDF_PATH.write_bytes(PDF_PATH.read_bytes())


def set_docx_font(run, size=None, bold=False, color=None):
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_docx_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text.upper())
    set_docx_font(run, 10, bold=True, color=(31, 78, 121))


def add_docx_entry(doc, entry):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(1)
    left_p.paragraph_format.line_spacing = Pt(11.5)
    r = left_p.add_run(entry["title"])
    set_docx_font(r, 9.4, bold=True)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(1)
    right_p.paragraph_format.line_spacing = Pt(11.5)
    r = right_p.add_run(entry.get("date", ""))
    set_docx_font(r, 8.7, color=(69, 81, 92))

    if entry.get("meta"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(10.8)
        r = p.add_run(entry["meta"])
        set_docx_font(r, 8.8, color=(69, 81, 92))

    for bullet in entry.get("bullets", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(10.8)
        r = p.add_run(bullet)
        set_docx_font(r, 8.6)


def make_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.52)
    sec.bottom_margin = Inches(0.56)
    sec.left_margin = Inches(0.58)
    sec.right_margin = Inches(0.58)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    name.paragraph_format.line_spacing = Pt(22)
    r = name.add_run(RESUME["name"])
    set_docx_font(r, 18, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    contact.paragraph_format.line_spacing = Pt(11)
    r = contact.add_run(" / ".join(RESUME["contact"]))
    set_docx_font(r, 8.5, color=(69, 81, 92))

    for section in RESUME["sections"]:
        if section.get("page_break_before"):
            doc.add_page_break()
        add_docx_heading(doc, section["title"])
        for entry in section["entries"]:
            add_docx_entry(doc, entry)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_pdf()
    make_docx()
    print(f"Wrote {PDF_PATH.relative_to(ROOT)}")
    print(f"Wrote {LEGACY_PDF_PATH.relative_to(ROOT)}")
    print(f"Wrote {DOCX_PATH.relative_to(ROOT)}")
